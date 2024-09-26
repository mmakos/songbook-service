import dataclasses
import json
import os
import sys

from docx import Document
from docx.text.paragraph import Paragraph

from authors import get_authors
from song_types import ICategory
from songbook_converter.song3 import parse_songs
from str_convert import title_to_unique_name

VERSION = "5.0.3_web"

categories = {
    "Kaczmarski & Gintrowski": "kaczmarski",
    "Kolędy": "carols",
    "Obozowe": "common",
    "Patriotyczne": "patriotic",
    "Religijne": "religious",
}

if __name__ == '__main__':
    if len(sys.argv) > 1:
        VERSION = sys.argv[1]

    doc = Document(os.path.abspath(f"docx/Śpiewnik-{VERSION}.docx"))
    pars = doc.paragraphs
    sections = dict()

    current_section = None
    for par in doc.paragraphs:
        par: Paragraph
        if par.style.name == "Heading 1":
            sections[par.text] = list()
            current_section = par.text
        elif current_section is not None:
            sections[current_section].append(par)

    with open("songbook_converter/songsWithoutAuthor.txt", "r", encoding="utf-8") as file:
        songs_without_author = [line.strip() for line in file.readlines()]

    authors = get_authors(f"docx/Śpiewnik-{VERSION}.htm")

    for section in sections:
        if section != 'Dodatki':
            songs = parse_songs(sections[section], authors)

            for song in songs:
                if song.performers is None and song.lyrics is None and song.composer is None and song.title not in songs_without_author:
                    print(f"Brak autora w piosence: {song.title}", file=sys.stderr)
                title = song.title.title()
                filename = title_to_unique_name(title)
                song.category = ICategory(categories[section], section)
                with open(f"songs/{filename}.json", "w", encoding="utf-8") as file:
                    file.write(json.dumps(
                        dataclasses.asdict(song, dict_factory=lambda x: {k: v for (k, v) in x if v is not None}),
                        ensure_ascii=False))
