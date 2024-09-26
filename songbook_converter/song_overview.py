from docx import Document
from docx.text.paragraph import Paragraph


def print_overviews(pars: list[Paragraph]):
    songs: list[tuple[str, list[Paragraph]]] = list()

    for par in pars:
        if par.style.name == "Heading 2" and len(par.text.strip()) > 0:
            songs.append((par.text.strip(), list()))
        elif len(songs) > 0:
            songs[-1][1].append(par)

    for song in songs:
        print_overview(song[0], song[1])


def print_overview(title: str, pars: list[Paragraph]):
    print(f"----- {title} -----")
    print("")
    for par in pars:
        print([run.text for run in par.runs])
    print("")
    print("")


VERSION = "5.0.3"

doc = Document(f"docx/Śpiewnik-{VERSION}.docx")
pars = doc.paragraphs
sections: dict[str, list[Paragraph]] = dict()

current_section = None
for par in doc.paragraphs:
    par: Paragraph
    if par.style.name == "Heading 1":
        sections[par.text] = list()
        current_section = par.text
    elif current_section is not None:
        sections[current_section].append(par)

for section in sections:
    if section != "Dodatki":
        print_overviews(sections[section])
